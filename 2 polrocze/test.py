def create_resume():
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        # Style settings
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Header
        name = doc.add_heading('MIKOŁAJ KUŚ', 0)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run('ul. Leśna 23/3, 76-150 Darłowo | Tel: 726 350 066\n')
        contact.add_run('Email: boydzburek@gmail.com | GitHub: github.com/Mi1ku')

        # Profile
        doc.add_heading('PROFIL ZAWODOWY', level=1)
        p = doc.add_paragraph('Ambitny uczeń 4. klasy technikum na kierunku Technik Programista ze zdaną kwalifikacją INF.03. ')
        p.add_run('Pasjonat technologii webowych i mobilnych. Posiadam praktyczne doświadczenie w administracji systemami CMS oraz zarządzaniu infrastrukturą webową. Obecnie rozwijam umiejętności z zakresu React oraz .NET MAUI, przygotowując się do egzaminu INF.04 oraz matury rozszerzonej z informatyki.')

        # Experience
        doc.add_heading('DOŚWIADCZENIE ZAWODOWE', level=1)
        exp = doc.add_paragraph()
        exp.add_run('Urząd Miejski w Darłowie (Referat Informatyczny)').bold = True
        exp.add_run('\nPraktykant Informatyk | Czerwiec 2025').italic = True
        
        doc.add_paragraph('Instalacja, pełna konfiguracja oraz bieżąca administracja systemem WordPress.', style='List Bullet')
        doc.add_paragraph('Zarządzanie plikami i bazami danych poprzez panel administracyjny cPanel.', style='List Bullet')
        doc.add_paragraph('Modyfikacja szablonów stron pod kątem responsywności (RWD) i estetyki.', style='List Bullet')
        doc.add_paragraph('Testowanie funkcjonalności i wydajności serwisów urzędowych.', style='List Bullet')

        # Skills
        doc.add_heading('UMIEJĘTNOŚCI TECHNICZNE', level=1)
        doc.add_paragraph('Web Development: React, HTML5, CSS3, JavaScript (podstawy), PHP (podstawy).', style='List Bullet')
        doc.add_paragraph('Programowanie: .NET MAUI, C# (podstawy), Python (podstawy), C++ (zakres rozszerzony).', style='List Bullet')
        doc.add_paragraph('Bazy danych & Narzędzia: SQL, WordPress, cPanel, Git (GitHub).', style='List Bullet')

        # Education
        doc.add_heading('WYKSZTAŁCENIE', level=1)
        doc.add_paragraph('2022 – obecnie: Zespół Szkół Agrotechnicznych w Sławnie | Technik Programista')
        doc.add_paragraph('2014 – 2022: Szkoła Podstawowa nr 3 w Darłowie')

        # Certificates
        doc.add_heading('CERTYFIKATY', level=1)
        doc.add_paragraph('Kwalifikacja zawodowa INF.03: Zdany egzamin państwowy.', style='List Bullet')
        doc.add_paragraph('Egzamin INF.04: W trakcie przygotowań (czerwiec 2026).', style='List Bullet')
        doc.add_paragraph('Ważna książeczka sanepidowska.', style='List Bullet')

        # RODO
        doc.add_paragraph('\n' + '-'*30)
        rodo = doc.add_paragraph()
        rodo_run = rodo.add_run('Wyrażam zgodę na przetwarzanie moich danych osobowych dla potrzeb niezbędnych do realizacji procesu rekrutacji (zgodnie z ustawą z dnia 10 maja 2018 roku o ochronie danych osobowych oraz zgodnie z RODO).')
        rodo_run.font.size = Pt(8)
        rodo_run.italic = True

        doc.save('CV_Mikolaj_Kus.docx')
        return "DOCX created successfully"
    except ImportError:
        # Fallback: create a plain text version so the script works without python-docx
        lines = [
            "MIKOŁAJ KUŚ",
            "=================",
            "",
            "ul. Leśna 23/3, 76-150 Darłowo | Tel: 726 350 066",
            "Email: boydzburek@gmail.com | GitHub: github.com/Mi1ku",
            "",
            "PROFIL ZAWODOWY",
            "Ambitny uczeń 4. klasy technikum na kierunku Technik Programista ze zdaną kwalifikacją INF.03.",
            "Pasjonat technologii webowych i mobilnych. Posiadam praktyczne doświadczenie w administracji systemami CMS oraz zarządzaniu infrastrukturą webową. Obecnie rozwijam umiejętności z zakresu React oraz .NET MAUI, przygotowując się do egzaminu INF.04 oraz matury rozszerzonej z informatyki.",
            "",
            "DOŚWIADCZENIE ZAWODOWE",
            "Urząd Miejski w Darłowie (Referat Informatyczny)",
            "Praktykant Informatyk | Czerwiec 2025",
            "- Instalacja, pełna konfiguracja oraz bieżąca administracja systemem WordPress.",
            "- Zarządzanie plikami i bazami danych poprzez panel administracyjny cPanel.",
            "- Modyfikacja szablonów stron pod kątem responsywności (RWD) i estetyki.",
            "- Testowanie funkcjonalności i wydajności serwisów urzędowych.",
            "",
            "UMIEJĘTNOŚCI TECHNICZNE",
            "- Web Development: React, HTML5, CSS3, JavaScript (podstawy), PHP (podstawy).",
            "- Programowanie: .NET MAUI, C# (podstawy), Python (podstawy), C++ (zakres rozszerzony).",
            "- Bazy danych & Narzędzia: SQL, WordPress, cPanel, Git (GitHub).",
            "",
            "WYKSZTAŁCENIE",
            "2022 – obecnie: Zespół Szkół Agrotechnicznych w Sławnie | Technik Programista",
            "2014 – 2022: Szkoła Podstawowa nr 3 w Darłowie",
            "",
            "CERTYFIKATY",
            "- Kwalifikacja zawodowa INF.03: Zdany egzamin państwowy.",
            "- Egzamin INF.04: W trakcie przygotowań (czerwiec 2026).",
            "- Ważna książeczka sanepidowska.",
            "",
            "-"*30,
            "Wyrażam zgodę na przetwarzanie moich danych osobowych dla potrzeb niezbędnych do realizacji procesu rekrutacji (zgodnie z ustawą z dnia 10 maja 2018 roku o ochronie danych osobowych oraz zgodnie z RODO)."
        ]
        with open('CV_Mikolaj_Kus.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        return "python-docx not available; TXT fallback created"
    except Exception as e:
        return f"Error: {e}"

if __name__ == '__main__':
    try:
        print(create_resume())
    except Exception as e:
        print(f"Error: {e}")